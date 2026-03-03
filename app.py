import streamlit as st
import cv2
import numpy as np
from ultralytics import YOLO
from PIL import Image
from docx import Document
from docx.shared import Inches
import io

# 1. Configuración de la App
st.set_page_config(page_title="FMS Analizador Pro", layout="centered")

if 'informe_clinico' not in st.session_state:
    st.session_state.informe_clinico = {}

st.sidebar.title("🩺 Gestión de Sesión")
nombre_paciente = st.sidebar.text_input("Nombre del Paciente", "Paciente Genérico")
test_seleccionado = st.sidebar.selectbox(
    "Seleccionar Test FMS",
    ["Deep Squat", "Hurdle Step", "Inline Lunge", "Shoulder Mobility"]
)

# Selector de lateralidad exclusivo para hombro
lado_hombro = ""
if test_seleccionado == "Shoulder Mobility":
    lado_hombro = st.sidebar.radio("Mano que sube (Lado):", ["Derecha", "Izquierda"])

if st.sidebar.button("🗑️ Nueva Sesión"):
    st.session_state.informe_clinico = {}
    st.rerun()

st.title(f"Evaluación: {test_seleccionado} {lado_hombro}")

@st.cache_resource
def load_model():
    return YOLO('yolov8n-pose.pt')

model = load_model()

# --- FUNCIONES DE APOYO ---
def get_angle(p1, p2):
    return np.degrees(np.arctan2(abs(p1[0]-p2[0]), abs(p1[1]-p2[1])))

def get_distance(p1, p2):
    return np.linalg.norm(p1 - p2)

def crear_word(datos_informe, nombre):
    doc = Document()
    doc.add_heading(f'Informe de Evaluación FMS', 0)
    doc.add_paragraph(f'Paciente: {nombre}')
    
    for test, datos in datos_informe.items():
        doc.add_heading(f'Test: {test}', level=1)
        doc.add_paragraph(f'Score: {datos["score"]}')
        doc.add_paragraph(f'Análisis Biomecánico: {datos["detalles"]}')
        doc.add_paragraph(f'Pautas Recomendadas: {datos["pautas"]}')
        
        img = Image.fromarray(datos["imagen"])
        img_stream = io.BytesIO()
        img.save(img_stream, format='PNG')
        img_stream.seek(0)
        
        doc.add_picture(img_stream, width=Inches(4))
        doc.add_paragraph("-" * 20)
        
    target_stream = io.BytesIO()
    doc.save(target_stream)
    target_stream.seek(0)
    return target_stream

# --- INTERFAZ DE CAPTURA ---
img_file_buffer = st.camera_input("Capturar Movimiento")

if img_file_buffer is not None:
    image = Image.open(img_file_buffer)
    frame = np.array(image)
    results = model(frame, verbose=False)
    annotated_frame = results[0].plot()
    kp = results[0].keypoints.xy[0].cpu().numpy()

    if len(kp) > 15:
        st.image(annotated_frame, caption="Captura Analizada")
        
        detalles, pautas, score = "", "", 1
        
        # --- LÓGICA DE CÁLCULO POR TEST ---
        if test_seleccionado == "Deep Squat":
            ang_torso = get_angle(kp[5], kp[11])
            ang_tibia = get_angle(kp[13], kp[15])
            dif = abs(ang_torso - ang_tibia)
            score = 3 if dif < 10 else 2 if dif < 25 else 1
            detalles = f"Torso: {ang_torso:.1f}°, Tibia: {ang_tibia:.1f}°. Diferencia: {dif:.1f}°."
            pautas = "Mejorar movilidad de tobillo y control lumbopélvico." if score < 3 else "Patrón óptimo."

        elif test_seleccionado == "Hurdle Step":
            verticalidad = get_angle(kp[5], kp[11])
            score = 3 if verticalidad < 5 else 2 if verticalidad < 15 else 1
            detalles = f"Inclinación anterior del tronco: {verticalidad:.1f}°."
            pautas = "Enfoque en estabilidad del core y fuerza de glúteo medio." if score < 3 else "Excelente estabilidad."

        elif test_seleccionado == "Inline Lunge":
            torso_lunge = get_angle(kp[5], kp[11])
            score = 3 if torso_lunge < 7 else 2 if torso_lunge < 20 else 1
            detalles = f"Verticalidad del torso: {torso_lunge:.1f}°."
            pautas = "Trabajar elasticidad de flexores de cadera y estabilidad lateral." if score < 3 else "Alineación correcta."

        elif test_seleccionado == "Shoulder Mobility":
            # Distancia entre manos (puntos 9 y 10)
            dist_manos = get_distance(kp[9], kp[10])
            # Estimación del tamaño de la mano (Distancia codo-muñeca / 3)
            tam_mano = get_distance(kp[7], kp[9]) / 2.8 
            
            proporcion = dist_manos / tam_mano
            score = 3 if proporcion < 1.0 else 2 if proporcion < 1.5 else 1
            detalles = f"Lado: {lado_hombro}. Distancia entre manos es {proporcion:.2f} veces el tamaño de la mano."
            pautas = "Mejorar movilidad torácica y flexibilidad de rotadores de hombro." if score < 3 else "Movilidad óptima."

        # Identificador único para guardar (evita que un lado sobreescriba al otro)
        id_test = f"{test_seleccionado} ({lado_hombro})" if lado_hombro else test_seleccionado

        if st.button(f"💾 Guardar {id_test}"):
            img_rgb = cv2.cvtColor(annotated_frame, cv2.COLOR_BGR2RGB)
            st.session_state.informe_clinico[id_test] = {
                "score": score, "detalles": detalles, "pautas": pautas, "imagen": img_rgb
            }
            st.toast(f"Test {id_test} guardado con éxito")

# --- SECCIÓN DE DESCARGA ---
if st.session_state.informe_clinico:
    st.divider()
    st.subheader("📋 Informe de Sesión")
    
    doc_word = crear_word(st.session_state.informe_clinico, nombre_paciente)
    
    st.download_button(
        label="📥 Descargar Informe Clínico en Word",
        data=doc_word,
        file_name=f"FMS_{nombre_paciente.replace(' ', '_')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    for t, d in st.session_state.informe_clinico.items():
        st.write(f"✅ **{t}** | Score: {d['score']}")
